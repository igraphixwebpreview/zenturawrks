#!/usr/bin/env python3
"""
Word Template to PDF Invoice Generator
Processes .docx templates with invoice data and converts to PDF
"""

import os
import json
import sys
import tempfile
from pathlib import Path
from typing import Dict, Any, List
from docx import Document
from docx.shared import Inches
from docx.table import Table
from docx2pdf import convert
import re

class InvoiceTemplateProcessor:
    def __init__(self, templates_dir="templates"):
        self.templates_dir = Path(templates_dir)
        self.templates_dir.mkdir(exist_ok=True)
        
    def replace_template_variables(self, text: str, data: Dict[str, Any]) -> str:
        """Replace template variables in text with actual data"""
        if not text:
            return ""
            
        # Replace {{variable}} patterns with actual data
        def replace_var(match):
            var_name = match.group(1)
            return str(data.get(var_name, f"{{{{{var_name}}}}}"))
        
        return re.sub(r'\{\{(\w+)\}\}', replace_var, text)
    
    def process_table_rows(self, table: Table, invoice_items: List[Dict[str, Any]], template_row_index: int = 1):
        """Process table rows for invoice items"""
        if not invoice_items or template_row_index >= len(table.rows):
            return
            
        # Get the template row
        template_row = table.rows[template_row_index]
        
        # Remove existing data rows (keep header and template)
        rows_to_remove = []
        for i in range(template_row_index + 1, len(table.rows)):
            rows_to_remove.append(i)
        
        # Remove from bottom to top to maintain indices
        for row_index in reversed(rows_to_remove):
            table._element.remove(table.rows[row_index]._element)
        
        # Add rows for each invoice item
        for item in invoice_items:
            # Create new row by copying template row
            new_row = table.add_row()
            
            # Copy cell content and replace variables
            for i, cell in enumerate(template_row.cells):
                if i < len(new_row.cells):
                    template_text = cell.text
                    new_text = self.replace_template_variables(template_text, {
                        'item_name': item.get('name', ''),
                        'item_description': item.get('description', ''),
                        'item_quantity': str(item.get('quantity', 1)),
                        'item_rate': f"${float(item.get('rate', 0)):.2f}",
                        'item_amount': f"${float(item.get('amount', 0)):.2f}"
                    })
                    new_row.cells[i].text = new_text
    
    def process_document(self, template_path: Path, invoice_data: Dict[str, Any]) -> Path:
        """Process a Word document template with invoice data"""
        
        # Open the template document
        doc = Document(template_path)
        
        # Prepare data for replacement
        replacement_data = {
            # Company information
            'company_name': invoice_data.get('companyName', 'Your Company'),
            'company_address': invoice_data.get('companyAddress', ''),
            'company_phone': invoice_data.get('companyPhone', ''),
            'company_email': invoice_data.get('companyEmail', ''),
            'company_website': invoice_data.get('companyWebsite', ''),
            
            # Invoice details
            'invoice_number': invoice_data.get('invoiceNumber', 'INV-001'),
            'invoice_date': invoice_data.get('invoiceDate', ''),
            'due_date': invoice_data.get('dueDate', ''),
            'status': invoice_data.get('status', 'draft'),
            
            # Client information
            'client_name': invoice_data.get('clientName', ''),
            'client_email': invoice_data.get('clientEmail', ''),
            'client_address': invoice_data.get('clientAddress', ''),
            'client_phone': invoice_data.get('clientPhone', ''),
            
            # Financial information
            'subtotal': f"${float(invoice_data.get('subtotal', 0)):.2f}",
            'discount_amount': f"${float(invoice_data.get('discountAmount', 0)):.2f}",
            'vat_rate': f"{float(invoice_data.get('vatRate', 0)):.1f}",
            'vat_amount': f"${float(invoice_data.get('vatAmount', 0)):.2f}",
            'deposit_amount': f"${float(invoice_data.get('depositAmount', 0)):.2f}",
            'total': f"${float(invoice_data.get('total', 0)):.2f}",
            'currency': '$',
            
            # Additional fields
            'notes': invoice_data.get('notes', ''),
            'payment_terms': invoice_data.get('paymentTerms', 'Net 30 days'),
            'thank_you_message': 'Thank you for your business!'
        }
        
        # Process paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text:
                paragraph.text = self.replace_template_variables(paragraph.text, replacement_data)
        
        # Process tables
        invoice_items = invoice_data.get('items', [])
        for table in doc.tables:
            # Check if this table contains item template codes
            has_item_templates = False
            template_row_index = -1
            
            for row_index, row in enumerate(table.rows):
                row_text = ' '.join(cell.text for cell in row.cells)
                if '{{item_' in row_text:
                    has_item_templates = True
                    template_row_index = row_index
                    break
            
            if has_item_templates and template_row_index >= 0:
                # This is an items table
                self.process_table_rows(table, invoice_items, template_row_index)
            else:
                # Regular table, just replace variables
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            cell.text = self.replace_template_variables(cell.text, replacement_data)
        
        # Process headers and footers
        for section in doc.sections:
            # Header
            if section.header:
                for paragraph in section.header.paragraphs:
                    if paragraph.text:
                        paragraph.text = self.replace_template_variables(paragraph.text, replacement_data)
            
            # Footer
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    if paragraph.text:
                        paragraph.text = self.replace_template_variables(paragraph.text, replacement_data)
        
        # Save processed document to temporary file
        temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_docx.name)
        temp_docx.close()
        
        return Path(temp_docx.name)
    
    def generate_invoice_pdf(self, invoice_data: Dict[str, Any], template_name: str = "default_invoice.docx") -> bytes:
        """Generate PDF invoice from template and data"""
        
        template_path = self.templates_dir / template_name
        
        # Check if template exists
        if not template_path.exists():
            raise FileNotFoundError(f"Template {template_name} not found in {self.templates_dir}")
        
        # Process the document
        processed_docx = self.process_document(template_path, invoice_data)
        
        try:
            # Convert to PDF
            temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
            temp_pdf.close()
            
            convert(str(processed_docx), temp_pdf.name)
            
            # Read PDF content
            with open(temp_pdf.name, 'rb') as f:
                pdf_content = f.read()
            
            return pdf_content
            
        finally:
            # Clean up temporary files
            if processed_docx.exists():
                processed_docx.unlink()
            if Path(temp_pdf.name).exists():
                Path(temp_pdf.name).unlink()

def main():
    """Command line interface for template processing"""
    try:
        # Read JSON input from stdin
        input_data = json.loads(sys.stdin.read())
        
        # Create processor
        processor = InvoiceTemplateProcessor()
        
        # Generate PDF
        template_name = input_data.get('template', 'default_invoice.docx')
        invoice_data = input_data.get('invoice_data', {})
        
        pdf_content = processor.generate_invoice_pdf(invoice_data, template_name)
        
        # Output base64 encoded PDF
        import base64
        pdf_base64 = base64.b64encode(pdf_content).decode('utf-8')
        
        result = {
            "success": True,
            "pdf_base64": pdf_base64,
            "message": "PDF generated successfully"
        }
        
        print(json.dumps(result))
        
    except Exception as e:
        error_result = {
            "success": False,
            "error": str(e),
            "message": "Failed to generate PDF"
        }
        print(json.dumps(error_result), file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()