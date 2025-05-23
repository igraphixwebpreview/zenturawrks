#!/usr/bin/env python3
"""
Create a default Word invoice template with professional styling
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

def create_default_invoice_template():
    """Create a professional invoice template with variables"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Company Header
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.columns[0].width = Inches(3.5)
    header_table.columns[1].width = Inches(3)
    
    # Company info (left side)
    company_cell = header_table.cell(0, 0)
    company_para = company_cell.paragraphs[0]
    company_run = company_para.add_run("{{companyName}}")
    company_run.font.size = Pt(18)
    company_run.font.bold = True
    
    company_cell.add_paragraph("{{companyAddress}}")
    company_cell.add_paragraph("Phone: {{companyPhone}}")
    company_cell.add_paragraph("Email: {{companyEmail}}")
    if "{{companyWebsite}}" != "":
        company_cell.add_paragraph("Website: {{companyWebsite}}")
    
    # Invoice title (right side)
    invoice_cell = header_table.cell(0, 1)
    invoice_para = invoice_cell.paragraphs[0]
    invoice_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    invoice_run = invoice_para.add_run("INVOICE")
    invoice_run.font.size = Pt(24)
    invoice_run.font.bold = True
    
    # Invoice details
    invoice_cell.add_paragraph(f"Invoice #: {{{{invoiceNumber}}}}")
    invoice_cell.add_paragraph(f"Date: {{{{invoiceDate}}}}")
    invoice_cell.add_paragraph(f"Due Date: {{{{dueDate}}}}")
    
    # Add some space
    doc.add_paragraph()
    
    # Bill To section
    bill_to_para = doc.add_paragraph()
    bill_to_run = bill_to_para.add_run("Bill To:")
    bill_to_run.font.size = Pt(14)
    bill_to_run.font.bold = True
    
    doc.add_paragraph("{{clientName}}")
    doc.add_paragraph("{{clientAddress}}")
    doc.add_paragraph("Email: {{clientEmail}}")
    
    # Add space before items table
    doc.add_paragraph()
    
    # Items table header
    items_table = doc.add_table(rows=1, cols=5)
    items_table.style = 'Table Grid'
    items_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    header_cells = items_table.rows[0].cells
    header_cells[0].text = "Item"
    header_cells[1].text = "Description"
    header_cells[2].text = "Rate"
    header_cells[3].text = "Qty"
    header_cells[4].text = "Amount"
    
    # Make header bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Template row for items (will be duplicated by template processor)
    row_cells = items_table.add_row().cells
    row_cells[0].text = "{{item.name}}"
    row_cells[1].text = "{{item.description}}"
    row_cells[2].text = "${{item.rate}}"
    row_cells[3].text = "{{item.quantity}}"
    row_cells[4].text = "${{item.amount}}"
    
    # Add space after table
    doc.add_paragraph()
    
    # Summary section
    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.autofit = False
    summary_table.columns[0].width = Inches(1.5)
    summary_table.columns[1].width = Inches(1)
    summary_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    
    # Summary rows
    summary_table.cell(0, 0).text = "Subtotal:"
    summary_table.cell(0, 1).text = "${{subtotal}}"
    
    summary_table.cell(1, 0).text = "Discount:"
    summary_table.cell(1, 1).text = "{{discount}}%"
    
    summary_table.cell(2, 0).text = "VAT:"
    summary_table.cell(2, 1).text = "{{vat}}%"
    
    # Total row (bold)
    total_label = summary_table.cell(3, 0)
    total_label.text = "Total:"
    total_amount = summary_table.cell(3, 1)
    total_amount.text = "${{total}}"
    
    # Make total row bold
    for cell in [total_label, total_amount]:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Notes section
    doc.add_paragraph()
    notes_para = doc.add_paragraph()
    notes_run = notes_para.add_run("Notes:")
    notes_run.font.bold = True
    doc.add_paragraph("{{notes}}")
    
    # Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph("Thank you for your business!")
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.runs[0]
    footer_run.font.italic = True
    
    # Save the template
    doc.save("templates/default_invoice.docx")
    print("Default invoice template created successfully!")

if __name__ == "__main__":
    create_default_invoice_template()